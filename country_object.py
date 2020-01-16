import sys
import os
import re
import collections
import json
import globals as g
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docxcompose.composer import Composer


class country_object(object):
    def __init__(self):
        self.country = ""
        self.rules_of_origin_scheme_sid = ""
        self.rules_of_origin_scheme_description = ""

    def write_JSON(self):
        # Get filenames
        path = os.path.join(os.getcwd(), "output")
        src_path = os.path.join(os.getcwd(), "source")
        folder_json = os.path.join(path, "_json")
        folder_docx = os.path.join(path, "_docx")
        folder_temp = os.path.join(path, "_temp")

        filename_json = "roo_" + self.rules_of_origin_scheme_description_formatted() + \
            ".json"
        filename_json = os.path.join(folder_json, filename_json)

        filename_docx = "roo_" + self.rules_of_origin_scheme_description_formatted() + \
            ".docx"
        filename_docx = os.path.join(folder_docx, filename_docx)

        filename_temp = os.path.join(folder_temp, "temp.docx")

        filename_template = os.path.join(src_path, "roo_template.docx")

        # Get the data
        sql = """SELECT heading, description, processing_rule, chapter, country, sequence
        FROM ml.roo_row where country = %s order by chapter, 6;"""
        params = [
            self.country
        ]
        cur = g.app.conn.cursor()
        cur.execute(sql, params)
        rows = cur.fetchall()

        master_object = collections.OrderedDict()
        chapters_object = []

        # Get the basic data from the database
        raw_list = []
        for row in rows:
            record_object = record_class(
                row[0], row[1], row[2], row[3], row[4], row[5])
            raw_list.append(record_object)

        # Assign to chapter objects
        chapters = []
        last_chapter = ""
        for raw_item in raw_list:
            r = row_class(raw_item.heading, raw_item.description,
                          raw_item.processing_rule)
            if raw_item.chapter != last_chapter:
                c = chapter_class(raw_item.chapter)
                chapters.append(c)
            c.rows.append(r)
            last_chapter = raw_item.chapter

        # Create JSON objects
        for c in chapters:
            chapter = collections.OrderedDict()
            chapter["number"] = c.number
            rows = []
            for item in c.rows:
                row = collections.OrderedDict()
                row["id"] = item.id
                row["description"] = item.description
                row["processing_rule"] = item.processing_rule
                rows.append(row)
            chapter["rows"] = rows

            chapters_object.append(chapter)

        # Append to the master object
        master_object["scheme_sid"] = self.rules_of_origin_scheme_sid
        master_object["scheme_description"] = self.rules_of_origin_scheme_description
        master_object["chapters"] = chapters_object

        j = json.dumps(master_object, indent=4)
        f = open(filename_json, "w+")
        f.write(j)
        f.close()

        write_word = True
        if write_word is False:
            return

        # Create the Word document
        document = Document()

        # Set the font
        font = document.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(10)

        # Set styles
        my_styles = document.styles
        """
        for st in my_styles:
            print(st.name)
        """

        p_style2 = my_styles.add_style('Small in table', WD_STYLE_TYPE.PARAGRAPH)
        p_style2.base_style = my_styles['Normal']
        p_style2.font.name = "Times New Roman"
        p_style2.font.size = Pt(9)
        p_style2.paragraph_format.space_before = Pt(5)
        p_style2.paragraph_format.space_after = Pt(5)

        title_style = my_styles['Title']
        title_style.font.name = "Times New Roman"
        title_style.font.size = Pt(24)
        title_style.font.color.rgb = RGBColor(0x00, 0x0c, 0x00)

        # Set the margins & orientation
        sections = document.sections
        margin = 1.5
        for section in sections:
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)

        document.add_heading("Rules of Origin for " + self.rules_of_origin_scheme_description, 0)

        if self.country in ("CA", "JP"):
            column_count = 2
        else:
            column_count = 3

        table = document.add_table(rows=1, cols=column_count)
        table.style = "Light List"
        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = 'Heading'
        hdr_cells[1].text = 'Description'
        if column_count > 2:
            hdr_cells[2].text = 'Processing rule'

        hdr_cells[0].paragraphs[0].style = p_style2
        hdr_cells[1].paragraphs[0].style = p_style2

        if column_count > 2:
            hdr_cells[2].paragraphs[0].style = p_style2
            hdr_cells[0].width = Inches(3)
            hdr_cells[1].width = Inches(6)
            hdr_cells[2].width = Inches(9)
        else:
            hdr_cells[0].width = Inches(3)
            hdr_cells[1].width = Inches(15)

        self.set_repeat_table_header(table.rows[0])

        for c in chapters:
            for item in c.rows:
                row_cells = table.add_row().cells

                self.display_html(row_cells[0], item.id, p_style2)
                row_cells[0].paragraphs[0].style = p_style2

                self.display_html(row_cells[1], item.description, p_style2)
                row_cells[1].paragraphs[0].style = p_style2

                if column_count > 2:
                    self.display_html(row_cells[2], item.processing_rule, p_style2)
                    row_cells[2].paragraphs[0].style = p_style2

        self.prevent_document_break(document)
        document.save(filename_temp)

        master = Document(filename_template)
        composer = Composer(master)
        doc_new = Document(filename_temp)
        composer.append(doc_new)
        composer.save(filename_docx)

    def display_html(self, rng, txt, sty):
        html_helper = HTMLHelper()
        par = rng.paragraphs[0]
        paras = txt.split("<br />")
        idx = 0
        for para in paras:
            if idx > 0:
                par = rng.add_paragraph()
                par.style = sty
            idx += 1
            run_map = html_helper.html_to_run_map(para)
            html_helper.insert_runs_from_html_map(par, run_map)
            if para[0:4] == "- - ":
                par.paragraph_format.first_line_indent = Pt(-10)
                par.paragraph_format.left_indent = Pt(10)
            elif para[0:2] == "- ":
                par.paragraph_format.first_line_indent = Pt(-5)
                par.paragraph_format.left_indent = Pt(5)

    def set_repeat_table_header(self, row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def rules_of_origin_scheme_description_formatted(self):
        s = self.rules_of_origin_scheme_description
        s = s.lower()
        s = s.replace(' - ', ' ')
        s = s.replace(' ', '_')
        s = s.replace(',', '')
        s = s.replace('(', '')
        s = s.replace(')', '')
        return s

    def prevent_document_break(self, document):
        tags = document.element.xpath('//w:tr')
        rows = len(tags)
        for row in range(0, rows):
            tag = tags[row]  # Specify which <w:r> tag you want
            child = OxmlElement('w:cantSplit')  # Create arbitrary tag
            tag.append(child)  # Append in the new tag


class record_class(object):
    def __init__(self, heading, description, processing_rule, chapter, country, sequence):
        self.heading = heading
        self.description = description
        self.processing_rule = processing_rule
        self.chapter = chapter
        self.country = country
        self.sequence = sequence


class chapter_class(object):
    def __init__(self, number):
        self.number = number
        self.rows = []


class row_class(object):
    def __init__(self, id, description, processing_rule):
        self.id = id
        self.description = description
        self.processing_rule = processing_rule


class HTMLHelper(object):
    """ Translates some html into word runs. """

    def __init__(self):
        self.get_tags = re.compile("(<[a-z,A-Z]+>)(.*?)(</[a-z,A-Z]+>)")

    def html_to_run_map(self, html_fragment):
        """ breakes an html fragment into a run map """
        ptr = 0
        run_map = []
        for match in self.get_tags.finditer(html_fragment):
            if match.start() > ptr:
                text = html_fragment[ptr:match.start()]
                if len(text) > 0:
                    run_map.append((text, "plain_text"))
            run_map.append((match.group(2), match.group(1)))
            ptr = match.end()
        if ptr < len(html_fragment) - 1:
            run_map.append((html_fragment[ptr:], "plain_text"))
        return run_map

    def insert_runs_from_html_map(self, paragraph, run_map):
        """ inserts some runs into a paragraph object. """
        for run_item in run_map:
            run = paragraph.add_run(run_item[0])
            if run_item[1] in ("<i>", "<em>"):
                run.italic = True
